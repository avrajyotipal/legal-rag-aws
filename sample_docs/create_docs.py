"""
Creates 5 realistic Indian civil case judgment documents as DOCX files.
Run: python sample_docs/create_docs.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).parent

def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def _para(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    if align:
        p.alignment = align
    return p

def _section(doc, title, body):
    _heading(doc, title, level=2)
    for para_text in body:
        _para(doc, para_text)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Property Partition Suit
# ══════════════════════════════════════════════════════════════════════════════
def create_doc1():
    doc = Document()
    doc.add_heading("IN THE COURT OF THE DISTRICT JUDGE, NEW DELHI", 0)
    _para(doc, "CIVIL SUIT NO. 245 OF 2019", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "RAMESH KUMAR S/O LATE SH. RATAN LAL, R/O 14-A, RAJA GARDEN, NEW DELHI – PLAINTIFF", bold=True)
    _para(doc, "VERSUS")
    _para(doc, "1. SURESH KUMAR S/O LATE SH. RATAN LAL, R/O 22, MODEL TOWN, NEW DELHI\n2. SMT. KAMLA DEVI W/O LATE SH. RATAN LAL, R/O 22, MODEL TOWN, NEW DELHI\n3. SMT. SAVITRI DEVI D/O LATE SH. RATAN LAL – DEFENDANTS", bold=True)
    doc.add_paragraph()

    _section(doc, "SUBJECT", [
        "Suit for Partition and Possession of ancestral property bearing Municipal No. 22, Model Town, New Delhi, admeasuring 250 sq. yards, along with the residential structure standing thereon (hereinafter referred to as the 'suit property')."
    ])

    _section(doc, "FACTS OF THE CASE", [
        "1. The suit property was owned by late Sh. Ratan Lal who died intestate on 12th March 2015, leaving behind three legal heirs: the plaintiff Ramesh Kumar (son), defendant no. 1 Suresh Kumar (son), and defendant no. 3 Smt. Savitri Devi (daughter). Defendant no. 2 Smt. Kamla Devi is the widow of the deceased.",
        "2. The deceased Sh. Ratan Lal had acquired the suit property in 1978 through a registered Sale Deed and had been residing therein along with his family. The property is assessed at a market value of approximately Rs. 2,50,00,000/- (Rupees Two Crore Fifty Lakhs only).",
        "3. After the demise of Sh. Ratan Lal, the defendants refused to effect a partition of the suit property and have been exclusively occupying the entire property despite repeated requests made by the plaintiff.",
        "4. The plaintiff claims 1/4th share in the suit property as per the provisions of the Hindu Succession Act, 1956, as amended by the Hindu Succession (Amendment) Act, 2005.",
        "5. A legal notice dated 15th August 2018 was served upon the defendants calling upon them to effect partition, but the same was neither complied with nor replied to.",
    ])

    _section(doc, "LEGAL ISSUES FRAMED", [
        "Issue 1: Whether the plaintiff is entitled to partition of the suit property?",
        "Issue 2: What is the share of each party in the suit property?",
        "Issue 3: Whether defendant no. 2 (widow) is entitled to any share?",
        "Issue 4: Whether the plaintiff is entitled to mesne profits?",
        "Issue 5: Relief.",
    ])

    _section(doc, "ARGUMENTS BY PLAINTIFF", [
        "Learned counsel for the plaintiff argued that under Section 8 of the Hindu Succession Act, 1956, the Class I heirs of the deceased are entitled to equal shares in the property. The four Class I heirs – widow (Kamla Devi) and three children (Ramesh, Suresh, and Savitri) – are each entitled to 1/4th share.",
        "Reliance was placed upon Uttam v. Saubhag Singh (2016) 4 SCC 68, wherein the Hon'ble Supreme Court held that the partition of property among Class I heirs under the Hindu Succession Act should be equal in all respects.",
        "The counsel further contended that the defendants have no right to exclusively occupy the property without effecting partition, and the plaintiff is entitled to mesne profits from the date of filing of the suit.",
    ])

    _section(doc, "ARGUMENTS BY DEFENDANTS", [
        "Learned counsel for the defendants argued that the suit property was not the exclusive self-acquired property of the deceased but formed part of a larger HUF (Hindu Undivided Family) property, and therefore the provisions of the Hindu Succession Act, 1956, as amended, would not apply in isolation.",
        "It was further contended that defendant no. 2, Smt. Kamla Devi, has been residing in the suit property since 1978 and has an absolute right of residence under Section 14 of the Hindu Succession Act, 1956.",
        "The defendants also raised a preliminary objection regarding the valuation of the suit for the purpose of court fees, submitting that the plaintiff had deliberately undervalued the suit.",
    ])

    _section(doc, "EVIDENCE ON RECORD", [
        "Ex. P-1: Death Certificate of Sh. Ratan Lal issued by Municipal Corporation of Delhi.",
        "Ex. P-2: Sale Deed dated 15th September 1978 in the name of Sh. Ratan Lal.",
        "Ex. P-3: Legal Notice dated 15th August 2018.",
        "Ex. P-4: Postal receipts and acknowledgment of service of notice.",
        "Ex. P-5: Certified copy of mutation entries from the revenue records.",
        "Ex. D-1: Alleged Will dated 01st January 2015 claimed by defendants (disputed).",
        "Ex. D-2: Electricity bills in the name of defendant no. 1.",
    ])

    _section(doc, "FINDINGS AND JUDGMENT", [
        "After careful consideration of the evidence on record and the arguments advanced by both the parties, this Court holds as under:",
        "On Issue 1 and 2: The suit property was the self-acquired property of late Sh. Ratan Lal, as is evident from the registered Sale Deed Ex. P-2. There is no credible evidence to establish the existence of any HUF. The alleged Will (Ex. D-1) is disbelieved as it has not been duly probated and the attesting witnesses have not supported its execution. Therefore, the property shall devolve by intestate succession under the Hindu Succession Act, 1956.",
        "This Court finds that all four Class I heirs – Smt. Kamla Devi (widow), Ramesh Kumar, Suresh Kumar, and Smt. Savitri Devi – are each entitled to 1/4th undivided share in the suit property.",
        "On Issue 3: Defendant no. 2 (widow) is entitled to a 1/4th share in her own right as a Class I heir under Section 8 read with Schedule of the Hindu Succession Act, 1956. Additionally, she has a right of residence under Section 14.",
        "On Issue 4: The plaintiff is entitled to mesne profits from the date of filing of the suit, i.e., from 01st April 2019, at the rate of Rs. 15,000/- per month, representing his proportionate share of the notional rent of the property.",
        "DECREE: A preliminary decree for partition is hereby passed. The suit property bearing Municipal No. 22, Model Town, New Delhi shall be partitioned into four equal shares of 1/4th each, one share to each of the four Class I heirs. The parties are directed to appear before the Local Commissioner to be appointed for drawing up the final decree.",
    ])

    _section(doc, "ORDER", [
        "CS No. 245/2019 is decreed in the above terms. The defendants are directed to pay the costs of the suit to the plaintiff. A final decree shall be drawn after the report of the Local Commissioner.",
        "Pronounced in open court on 14th February 2024.",
        "Sd/- DISTRICT JUDGE, NORTH WEST, NEW DELHI",
    ])

    out = OUT_DIR / "01_property_partition_suit_delhi.docx"
    doc.save(out)
    print(f"Created: {out}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: Commercial Contract Dispute
# ══════════════════════════════════════════════════════════════════════════════
def create_doc2():
    doc = Document()
    doc.add_heading("IN THE HIGH COURT OF DELHI AT NEW DELHI", 0)
    _para(doc, "O.M.P. (COMM) 187 OF 2021", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "(Under Section 34 of the Arbitration and Conciliation Act, 1996)", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "M/s ABC CONSTRUCTIONS PVT. LTD.\nRegistered Office: Plot No. 5, Okhla Industrial Area, Phase-II, New Delhi – 110020\n... PETITIONER", bold=True)
    _para(doc, "VERSUS")
    _para(doc, "M/s XYZ DEVELOPERS LTD.\nRegistered Office: Tower A, DLF Cyber City, Gurugram, Haryana – 122002\n... RESPONDENT", bold=True)
    doc.add_paragraph()

    _section(doc, "SUBJECT", [
        "Petition under Section 34 of the Arbitration and Conciliation Act, 1996 for setting aside the Arbitral Award dated 30th June 2021, passed in arbitration proceedings arising out of a Construction Agreement dated 01st April 2018 between the parties."
    ])

    _section(doc, "BACKGROUND AND FACTS", [
        "1. The petitioner, M/s ABC Constructions Pvt. Ltd., entered into a Construction Agreement dated 01st April 2018 with the respondent M/s XYZ Developers Ltd. for the construction of a commercial complex at Sector 62, Noida, Uttar Pradesh, for a contract sum of Rs. 45,00,00,000/- (Rupees Forty-Five Crores only).",
        "2. The Agreement stipulated completion of the project within 24 months from the date of handing over of the site, i.e., by 31st March 2020, subject to force majeure and extension provisions.",
        "3. The petitioner claims that the respondent failed to hand over the site in a timely manner, caused repeated design changes, and failed to make payments as per the agreed payment schedule, resulting in delay and additional costs to the tune of Rs. 8,75,00,000/-.",
        "4. The respondent invoked the arbitration clause (Clause 26 of the Agreement) alleging that the petitioner had abandoned the project midway and was liable for liquidated damages of Rs. 5,00,00,000/- at the rate of 0.5% per week of the contract value for the delay of 22 weeks.",
        "5. The Sole Arbitrator, retired Justice M.K. Sharma (Retd. Judge, Delhi High Court), passed an award dated 30th June 2021, rejecting the petitioner's claims in entirety and awarding liquidated damages of Rs. 3,00,00,000/- in favour of the respondent.",
    ])

    _section(doc, "GROUNDS FOR CHALLENGE UNDER SECTION 34", [
        "Ground 1 – Patent Illegality: The Arbitrator failed to consider the force majeure clause (Clause 18) which excused the petitioner from liability for delays caused by acts of the employer (respondent).",
        "Ground 2 – Violation of Natural Justice: Key documents submitted by the petitioner, specifically the site handover correspondence and design change notices, were not considered by the Arbitrator.",
        "Ground 3 – Conflict with Public Policy: The Arbitrator applied liquidated damages without examining whether the same were a genuine pre-estimate of loss as required by Section 74 of the Indian Contract Act, 1872.",
        "Ground 4 – Perversity: The finding that the petitioner abandoned the project is perverse and contrary to the evidence on record.",
    ])

    _section(doc, "RESPONDENT'S COUNTER CONTENTIONS", [
        "1. The scope of challenge under Section 34 of the Act is narrow and limited to the grounds specified therein. Mere errors of law or fact do not qualify as grounds for setting aside an award.",
        "2. The learned Arbitrator has considered all documents and his findings on abandonment are based on evidence. The Court cannot re-appreciate evidence under Section 34.",
        "3. Reliance placed on Associate Builders v. Delhi Development Authority (2015) 3 SCC 49, wherein the Supreme Court held that the Court's jurisdiction under Section 34 is supervisory, not appellate.",
        "4. The liquidated damages clause was a mutual agreement and the respondent suffered real commercial losses due to the delay in project completion.",
    ])

    _section(doc, "APPLICABLE LAW", [
        "Section 34, Arbitration and Conciliation Act, 1996 – Grounds for setting aside an arbitral award.",
        "Section 74, Indian Contract Act, 1872 – Compensation for breach of contract where penalty is stipulated.",
        "Ssangyong Engineering v. NHAI (2019) 15 SCC 131 – The award must not be so perverse or so irrational that no reasonable person could have arrived at it.",
        "ONGC Ltd. v. Saw Pipes Ltd. (2003) 5 SCC 705 – Patent illegality on the face of the award.",
    ])

    _section(doc, "COURT'S ANALYSIS AND DECISION", [
        "After examining the arbitral record, the Award, and the submissions of the parties, this Court finds as follows:",
        "On Ground 1 (Force Majeure): The Arbitrator did address the force majeure clause at paragraphs 45-52 of the Award. The finding that the delays were not attributable to force majeure events is a plausible view based on the evidence. This Court cannot substitute its view under Section 34.",
        "On Ground 2 (Natural Justice): The Arbitrator has specifically mentioned the site handover correspondence at paragraph 38 of the Award and offered reasons for not relying upon the same. No violation of natural justice is made out.",
        "On Ground 3 (Liquidated Damages): This ground has merit. The Arbitrator failed to examine whether Rs. 5,00,00,000/- represented a genuine pre-estimate of loss. The mere invocation of a liquidated damages clause does not dispense with the requirement under Section 74 of the Contract Act to establish actual loss. The award of Rs. 3,00,00,000/- in liquidated damages is set aside.",
        "On Ground 4 (Perversity): The finding of abandonment is based on site inspection reports and correspondence. The same cannot be characterized as perverse.",
        "DECISION: The petition is partly allowed. The Arbitral Award dated 30th June 2021 is set aside to the limited extent of the award of liquidated damages of Rs. 3,00,00,000/-. The matter is remitted to the Arbitral Tribunal for fresh adjudication on the question of liquidated damages. The award rejecting the petitioner's claims is upheld.",
        "Pronounced on: 22nd November 2023.\nSd/- HON'BLE JUSTICE PRADEEP NANDA, DELHI HIGH COURT",
    ])

    out = OUT_DIR / "02_commercial_arbitration_delhi_hc.docx"
    doc.save(out)
    print(f"Created: {out}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: Matrimonial Dispute – Divorce and Maintenance
# ══════════════════════════════════════════════════════════════════════════════
def create_doc3():
    doc = Document()
    doc.add_heading("IN THE FAMILY COURT, BANDRA, MUMBAI", 0)
    _para(doc, "MATRIMONIAL PETITION NO. 892 OF 2020", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "SMT. PRIYA SHARMA W/O RAJESH SHARMA, AGE: 34 YEARS,\nR/O FLAT NO. 4B, SHANTI NAGAR CHS, ANDHERI (WEST), MUMBAI – 400058\n... PETITIONER", bold=True)
    _para(doc, "VERSUS")
    _para(doc, "RAJESH SHARMA S/O SH. VINOD SHARMA, AGE: 38 YEARS,\nR/O 12, LOKHANDWALA COMPLEX, ANDHERI (WEST), MUMBAI – 400053\n... RESPONDENT", bold=True)
    doc.add_paragraph()

    _section(doc, "NATURE OF PETITION", [
        "Petition for dissolution of marriage by a decree of divorce under Section 13(1)(ia) and (ib) of the Hindu Marriage Act, 1955, on the grounds of cruelty and desertion, along with a claim for permanent alimony under Section 25 and custody of the minor child under Section 26 of the Hindu Marriage Act, 1955."
    ])

    _section(doc, "BACKGROUND", [
        "1. The petitioner Smt. Priya Sharma (née Verma) and the respondent Rajesh Sharma were married on 20th November 2012 at Mumbai, according to Hindu rites and customs. A child, Master Aryan Sharma, was born from the wedlock on 14th June 2015, who is presently aged 9 years.",
        "2. The couple lived together at the matrimonial home in Andheri (West), Mumbai. The petitioner alleges that within 6 months of marriage, the respondent started subjecting her to physical and mental cruelty, making demands for additional dowry, and abusing her in the presence of her in-laws.",
        "3. On 03rd March 2019, the respondent allegedly assaulted the petitioner, as a result of which she was hospitalised for 3 days at Kokilaben Dhirubhai Ambani Hospital. The medical records have been produced as Ex. P-7.",
        "4. The respondent left the matrimonial home in September 2019 and has neither contacted the petitioner nor maintained the petitioner and child since then, constituting desertion for a period of more than two years.",
        "5. The respondent is employed as a Senior Manager in a multinational company and draws a salary of approximately Rs. 2,00,000/- per month.",
    ])

    _section(doc, "RELIEFS CLAIMED", [
        "a) A decree of divorce under Section 13(1)(ia) – cruelty and Section 13(1)(ib) – desertion.",
        "b) Permanent alimony of Rs. 50,000/- per month under Section 25, HMA.",
        "c) Exclusive custody of Master Aryan Sharma (minor child), with reasonable visitation rights to respondent under Section 26, HMA.",
        "d) Litigation expenses.",
    ])

    _section(doc, "RESPONDENT'S DEFENSE", [
        "1. The respondent denied all allegations of cruelty and desertion. He submitted that it was the petitioner herself who left the matrimonial home after creating a scene.",
        "2. The respondent contended that he is maintaining the child and making monthly payments of Rs. 15,000/- to the petitioner, though the petitioner denies receipt of the same.",
        "3. On the question of alimony, the respondent submitted that the petitioner is a post-graduate and is capable of earning, and therefore permanent alimony should not be granted.",
        "4. Regarding custody, the respondent argued that the child should be in joint custody and the petitioner should not be granted exclusive custody.",
    ])

    _section(doc, "LAW APPLIED", [
        "Section 13(1)(ia), Hindu Marriage Act, 1955 – Cruelty as a ground for divorce.",
        "Section 13(1)(ib), Hindu Marriage Act, 1955 – Desertion for 2 years as a ground for divorce.",
        "Section 25, Hindu Marriage Act, 1955 – Permanent alimony and maintenance.",
        "Section 26, Hindu Marriage Act, 1955 – Custody of children.",
        "Naveen Kohli v. Neelu Kohli (2006) 4 SCC 558 – Test for cruelty in matrimonial cases.",
        "Savitaben Somabhai Bhatiya v. State of Gujarat (2005) 3 SCC 636 – Mental cruelty includes conduct causing reasonable apprehension.",
    ])

    _section(doc, "FINDINGS", [
        "On Cruelty (Section 13(1)(ia)): The medical records (Ex. P-7) corroborated by the testimony of the treating physician (PW-3) establish beyond reasonable doubt that the petitioner suffered physical injury at the hands of the respondent. Mental cruelty through continuous demand for dowry is also proved through the testimony of the petitioner's parents (PW-1, PW-2). Cruelty is established.",
        "On Desertion (Section 13(1)(ib)): The respondent left the matrimonial home in September 2019. The period of desertion exceeds two years. The animus deserendi (intention to desert) is established from the respondent's conduct and the WhatsApp messages (Ex. P-9). Desertion is proved.",
        "On Alimony (Section 25): Considering the respondent's salary of Rs. 2,00,000/- per month and the petitioner's lack of current employment, permanent alimony of Rs. 35,000/- per month is awarded.",
        "On Custody (Section 26): Exclusive custody of Master Aryan Sharma is awarded to the petitioner. The respondent shall have visitation rights every alternate weekend from Friday 6 PM to Sunday 6 PM and during half of school vacation periods.",
    ])

    _section(doc, "DECREE", [
        "1. The marriage solemnised on 20th November 2012 between Smt. Priya Sharma and Sh. Rajesh Sharma is hereby dissolved by a decree of divorce.",
        "2. The respondent shall pay permanent alimony of Rs. 35,000/- per month to the petitioner, payable by the 5th of every month.",
        "3. Custody of Master Aryan Sharma is awarded to the petitioner with visitation rights to the respondent as specified above.",
        "4. The respondent shall pay litigation costs of Rs. 25,000/-.",
        "Pronounced on: 10th January 2024.\nSd/- PRINCIPAL JUDGE, FAMILY COURT, BANDRA, MUMBAI",
    ])

    out = OUT_DIR / "03_matrimonial_divorce_mumbai_family_court.docx"
    doc.save(out)
    print(f"Created: {out}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: Tenancy / Eviction Dispute
# ══════════════════════════════════════════════════════════════════════════════
def create_doc4():
    doc = Document()
    doc.add_heading("BEFORE THE COURT OF SMALL CAUSES, MUMBAI", 0)
    _para(doc, "R.A.D. SUIT NO. 1124 OF 2018", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "(Under Section 33 of the Maharashtra Rent Control Act, 1999)", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "M/s SHAH BROTHERS\nThrough its partner: Sh. Hiren Shah,\n45, Linking Road, Bandra (West), Mumbai – 400050\n... PLAINTIFF / LANDLORD", bold=True)
    _para(doc, "VERSUS")
    _para(doc, "M/s GOPAL ENTERPRISES\nThrough its proprietor: Sh. Gopal Krishnamurthy,\n45, Linking Road, Bandra (West), Mumbai – 400050 (same premises)\n... DEFENDANT / TENANT", bold=True)
    doc.add_paragraph()

    _section(doc, "SUBJECT MATTER", [
        "Suit for eviction of the defendant/tenant from the commercial premises on the ground floor of the property bearing House No. 45, Linking Road, Bandra (West), Mumbai (hereinafter 'the suit premises'), on grounds of (a) non-payment of rent and (b) subletting without consent, under Sections 15 and 16 of the Maharashtra Rent Control Act, 1999."
    ])

    _section(doc, "FACTS", [
        "1. The plaintiff M/s Shah Brothers is the owner of the property bearing House No. 45, Linking Road, Bandra (West), Mumbai. The defendant M/s Gopal Enterprises has been a tenant of the ground floor commercial premises since 01st April 2005, on a monthly rental of Rs. 25,000/-.",
        "2. The monthly rent was last revised to Rs. 35,000/- per month vide a rent agreement dated 01st April 2015 for a period of 3 years (2015–2018). After expiry of the agreement, the defendant continued in possession as a statutory tenant under the Maharashtra Rent Control Act, 1999.",
        "3. The plaintiff alleges that the defendant has failed to pay rent since April 2017, resulting in an arrear of rent of Rs. 8,75,000/- (35 months × Rs. 25,000/-) as on February 2020.",
        "4. Furthermore, the plaintiff discovered that the defendant has sublet a portion of the suit premises to a third party, M/s Patel Trading Co., without the plaintiff's written consent, in violation of Section 15 of the MRC Act, 1999.",
        "5. A notice under Section 16(2) of the MRC Act was served upon the defendant on 15th November 2017, but the defendant failed to vacate or pay the arrears.",
    ])

    _section(doc, "DEFENDANT'S DEFENSE", [
        "1. The defendant denied the allegation of non-payment, contending that rent was paid in cash and receipts were not issued by the plaintiff. The defendant claimed that the plaintiff deliberately refused to accept rent.",
        "2. On the subletting ground, the defendant submitted that M/s Patel Trading Co. was not a sub-tenant but a business associate who occasionally used the premises, and there was no formal subletting arrangement.",
        "3. The defendant also challenged the maintainability of the suit, contending that the plaintiff had not complied with the pre-requisite notice under Section 16(2) of the MRC Act before filing the suit.",
        "4. Reliance was placed upon Keku Balsara v. Adi B. Billimoria AIR 1959 Bom 119, for the proposition that a landlord cannot refuse rent and then claim non-payment.",
    ])

    _section(doc, "EVIDENCE", [
        "PW-1: Sh. Hiren Shah, partner of plaintiff, who testified to the tenancy, rent arrears, and subletting.",
        "PW-2: Sh. Ramesh Mehta, an employee of the plaintiff, who identified M/s Patel Trading Co. persons entering the suit premises regularly.",
        "Ex. P-1: Original Rent Agreement dated 01st April 2015.",
        "Ex. P-2: Rent receipts from 2005 to March 2017.",
        "Ex. P-3: Photographs showing signboard of 'M/s Patel Trading Co.' inside the suit premises.",
        "Ex. P-4: Notice under Section 16(2) MRC Act dated 15th November 2017.",
        "DW-1: Sh. Gopal Krishnamurthy, defendant, denying all allegations.",
        "Ex. D-1: Bank statements purportedly showing cash withdrawals (not directly evidencing payment of rent).",
    ])

    _section(doc, "LEGAL ANALYSIS", [
        "On Non-Payment of Rent: Under Section 16(1)(a) of the MRC Act, 1999, non-payment of rent is a valid ground for eviction. The burden lies on the tenant to prove payment. The defendant has failed to produce any valid receipt or other reliable evidence of payment. The bank statements (Ex. D-1) show cash withdrawals but do not establish payment to the plaintiff. The allegation of deliberate refusal to accept rent is not substantiated by any evidence such as a money order or deposit in court. Non-payment of rent is established.",
        "On Subletting: Section 15 of the MRC Act prohibits a tenant from subletting, assigning, or parting with possession of the premises or any part thereof without prior written consent of the landlord. Ex. P-3 (photographs) and the testimony of PW-2 establish that a third party, M/s Patel Trading Co., was in actual occupation and running a business from the suit premises. The defendant's explanation that they were merely 'business associates' is not credible. Unlawful subletting is established.",
        "On Notice: Ex. P-4 is a valid notice under Section 16(2) of the MRC Act. The objection to maintainability is overruled.",
    ])

    _section(doc, "JUDGMENT AND DECREE", [
        "The R.A.D. Suit No. 1124 of 2018 is decreed in favour of the plaintiff. The defendant M/s Gopal Enterprises, its proprietor, servants, agents and sub-tenants are directed to hand over vacant and peaceful possession of the suit premises at 45, Linking Road, Bandra (West), Mumbai, to the plaintiff M/s Shah Brothers within 30 days of this judgment.",
        "The defendant is directed to pay rent arrears of Rs. 8,75,000/- along with interest at 12% per annum from April 2017 till date of realisation.",
        "The defendant is directed to pay litigation costs of Rs. 30,000/-.",
        "Pronounced on: 18th March 2024.\nSd/- JUDGE, COURT OF SMALL CAUSES, MUMBAI",
    ])

    out = OUT_DIR / "04_tenancy_eviction_mumbai_small_causes.docx"
    doc.save(out)
    print(f"Created: {out}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 5: Medical Negligence / Consumer Case
# ══════════════════════════════════════════════════════════════════════════════
def create_doc5():
    doc = Document()
    doc.add_heading("BEFORE THE DISTRICT CONSUMER DISPUTES REDRESSAL COMMISSION, CHENNAI – I", 0)
    _para(doc, "CONSUMER COMPLAINT NO. 412 OF 2022", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "(Under Section 35 of the Consumer Protection Act, 2019)", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "DR. ANJALI MEHTA D/O SH. RAVI MEHTA, AGE: 42 YEARS,\nR/O FLAT 7, GREEN PARK APARTMENTS, VELACHERY, CHENNAI – 600042\n... COMPLAINANT", bold=True)
    _para(doc, "VERSUS")
    _para(doc, "1. CITY MULTISPECIALTY HOSPITAL PVT. LTD.\n   RA Road, Teynampet, Chennai – 600018\n2. DR. S. KRISHNASWAMY, M.S. (ORTHO),\n   Consultant Orthopaedic Surgeon, City Multispecialty Hospital\n... OPPOSITE PARTIES", bold=True)
    doc.add_paragraph()

    _section(doc, "NATURE OF COMPLAINT", [
        "The complainant alleges medical negligence and deficiency in service on the part of the opposite parties during and after a knee replacement surgery (Total Knee Arthroplasty – TKA) performed at Opposite Party No. 1 hospital on 15th March 2021, resulting in permanent disability and causing physical, mental, and financial suffering."
    ])

    _section(doc, "COMPLAINT AVERMENTS", [
        "1. The complainant, Dr. Anjali Mehta, a practising dentist, was suffering from severe osteoarthritis of the right knee for several years. After consultation, OP No. 2 (Dr. Krishnaswamy) advised Total Knee Arthroplasty (TKA) of the right knee.",
        "2. The surgery was performed on 15th March 2021 at OP No. 1 hospital. Total charges paid were Rs. 4,50,000/- (Rupees Four Lakhs Fifty Thousand only).",
        "3. Post-surgery, the complainant developed severe infection (septic arthritis) at the surgical site. Despite multiple complaints, the treating team attributed the fever and swelling to post-operative inflammation and did not conduct adequate investigation.",
        "4. By June 2021, the infection had spread to the bone (osteomyelitis), necessitating a revision surgery and removal of the prosthetic implant, leaving the complainant with a fused joint and permanent disability.",
        "5. The complainant had to undergo treatment at CMC Vellore, spending an additional Rs. 8,20,000/-. The complainant, being a dentist, lost approximately Rs. 15,00,000/- in professional income during the period of disability.",
        "6. Expert opinion from Dr. P. Raghunathan, Professor of Orthopaedics, Madras Medical College, was obtained, opining that the infection could have been prevented had proper sterile protocols been followed during surgery, and that the post-operative management fell short of established medical standards.",
    ])

    _section(doc, "OPPOSITE PARTIES' REPLY", [
        "1. The OPs denied any negligence, stating that the surgery was performed as per standard protocols and with full informed consent.",
        "2. Post-operative infection is a known complication of TKA with an incidence of 1–2% globally and is not indicative of negligence per se.",
        "3. As soon as the infection was diagnosed (culture report dated 28th April 2021), appropriate antibiotics were started and the complainant was managed as per IDSA guidelines.",
        "4. The OPs questioned the credentials and independence of the expert Dr. Raghunathan, alleging that he was a professional associate of the complainant's family.",
    ])

    _section(doc, "APPLICABLE LEGAL STANDARDS", [
        "Section 2(42), Consumer Protection Act, 2019 – Definition of 'service'.",
        "Section 2(11), Consumer Protection Act, 2019 – Definition of 'deficiency'.",
        "Jacob Mathew v. State of Punjab (2005) 6 SCC 1 – The standard of care in medical negligence cases is that of a reasonably competent professional, not the highest possible standard.",
        "Martin F. D'Souza v. Mohd. Ishfaq (2009) 3 SCC 1 – Doctors are not insurers; medical negligence requires proof of lack of reasonable care.",
        "Spring Meadows Hospital v. Harjol Ahluwalia (1998) 4 SCC 39 – Consumer fora have jurisdiction over medical negligence cases.",
    ])

    _section(doc, "ANALYSIS BY COMMISSION", [
        "1. The Commission carefully examined the hospital records, the operative notes, culture reports, and the expert opinion.",
        "2. The culture report dated 28th April 2021 revealed Staphylococcus aureus infection, a common indicator of breach of sterile technique during surgery. The hospital's Infection Control Audit for March 2021, obtained under Right to Information application, showed that the OT where the surgery was performed had recorded two other post-operative infections in the same month.",
        "3. The Commission notes that the post-operative management was delayed. The fever chart showed consistent elevated temperature from 22nd March 2021, yet a culture report was ordered only on 24th April 2021 – a gap of 33 days. This delay constitutes deficiency in service.",
        "4. Regarding the expert opinion, the Commission accepts the opinion of Dr. Raghunathan as credible and impartial. The OPs have not produced any counter-expert opinion.",
        "5. On the question of res ipsa loquitur – while it cannot be applied mechanically in medical cases, the cluster of infections in the same OT in March 2021 raises a strong inference of systemic failure in infection control protocols.",
    ])

    _section(doc, "AWARD", [
        "The Consumer Complaint No. 412 of 2022 is allowed. The Opposite Parties are held jointly and severally liable for medical negligence and deficiency in service.",
        "1. Compensation for medical expenses: Rs. 8,20,000/- (treatment at CMC Vellore and other expenses).",
        "2. Compensation for loss of professional income: Rs. 6,00,000/- (assessed at a conservative estimate).",
        "3. Compensation for pain, suffering, and permanent disability: Rs. 10,00,000/-.",
        "4. Compensation for mental agony and harassment: Rs. 2,00,000/-.",
        "5. Litigation costs: Rs. 50,000/-.",
        "Total Award: Rs. 26,70,000/- (Rupees Twenty-Six Lakhs Seventy Thousand only), payable within 45 days of this order, failing which the amount shall carry interest at 9% per annum.",
        "Pronounced on: 05th March 2024.\nSd/- PRESIDENT, DISTRICT CONSUMER DISPUTES REDRESSAL COMMISSION, CHENNAI – I",
    ])

    out = OUT_DIR / "05_medical_negligence_consumer_commission_chennai.docx"
    doc.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    OUT_DIR.mkdir(exist_ok=True)
    print("Creating 5 Indian civil case documents...\n")
    create_doc1()
    create_doc2()
    create_doc3()
    create_doc4()
    create_doc5()
    print("\nAll 5 documents created successfully in sample_docs/")
